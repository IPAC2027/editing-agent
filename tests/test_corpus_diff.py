"""Extracting the editors' real corrections from a before/after pair.

Most of these tests are about what must *not* appear. A diff that reports a
rewrapped paragraph as forty corrections is worse than no diff: it buries the
real edits and it inflates every count computed from them.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from src.corpus.diff import (
    diff_paper,
    diff_paragraph,
    diff_paragraphs,
    latex_paragraphs,
    word_paragraphs,
    write,
)
from src.corpus.index import load


def _hunks(before: str, after: str):
    return diff_paragraphs(latex_paragraphs(before), latex_paragraphs(after),
                           paper="X", source="latex")


# ---------------------------------------------------------------------------
# Noise that must not be reported
# ---------------------------------------------------------------------------

def test_rewrapping_a_paragraph_is_not_a_correction():
    """The single biggest source of noise in a LaTeX diff."""
    before = "The gun runs at 5 MHz\nwith a 250 pC bunch and a\nlong tail."
    after = "The gun runs at 5 MHz with a 250 pC bunch and a long tail."

    assert _hunks(before, after) == []


def test_trailing_whitespace_is_not_a_correction():
    assert _hunks("A line.   ", "A line.") == []


def test_a_non_breaking_space_IS_a_correction():
    """The distinction the whitespace filter exists to preserve.

    ``10 MeV`` -> ``10~MeV`` is the most common real correction in the corpus;
    a filter that treated it as whitespace would delete the evidence for the
    agent's most-used check.
    """
    hunks = _hunks("a beam of 10 MeV protons", "a beam of 10~MeV protons")

    assert [(h.before, h.after) for h in hunks] == [("10 MeV", "10~MeV")]


# ---------------------------------------------------------------------------
# Granularity
# ---------------------------------------------------------------------------

def test_a_correction_is_the_words_that_changed_not_the_paragraph():
    hunks = diff_paragraph("The gun runs at 5 mhz today", "The gun runs at 5 MHz today",
                           paper="X", unit=0, source="latex")

    assert len(hunks) == 1
    assert (hunks[0].before, hunks[0].after) == ("mhz", "MHz")
    assert hunks[0].words == 1


def test_several_corrections_in_one_paragraph_are_separate_hunks():
    hunks = _hunks(r"5 mhz with 250 pC \cite{a}\cite{b}.",
                   r"5 MHz with 250~pC \cite{a, b}.")

    assert len(hunks) == 3
    assert {h.before for h in hunks} == {"mhz", "250 pC", r"\cite{a}\cite{b}."}


def test_context_is_kept_so_a_hunk_can_be_judged_later():
    hunks = _hunks("The beam energy was 10 MeV at the target.",
                   "The beam energy was 10~MeV at the target.")

    assert "beam energy" in hunks[0].context


def test_a_wholesale_rewrite_is_recorded_but_marked_large():
    """A rewritten paragraph is not one correction and must not be counted as one.

    Below the similarity threshold the two are different paragraphs, so the
    whole block is kept as a single hunk and flagged, rather than exploded into
    a hundred word-level 'corrections' that would swamp every count.
    """
    before = " ".join(f"alpha{i} bravo charlie" for i in range(20))
    after = " ".join(f"zulu{i} yankee xray" for i in range(20))

    hunks = _hunks(before, after)

    assert len(hunks) == 1
    assert hunks[0].large is True
    assert hunks[0].words > 25


def test_many_independent_word_fixes_are_many_hunks():
    """The other direction: a paragraph an editor corrected in six places is
    six corrections, and collapsing them would lose five."""
    before = "we saw 1 mhz and 2 mhz and 3 mhz in the ring"
    after = "we saw 1 MHz and 2 MHz and 3 MHz in the ring"

    hunks = _hunks(before, after)

    assert [h.before for h in hunks] == ["mhz", "mhz", "mhz"]


def test_an_inserted_paragraph_is_one_hunk_not_a_word_diff_against_nothing():
    hunks = _hunks("First.\n\nThird.", "First.\n\nSecond one here.\n\nThird.")

    inserts = [h for h in hunks if h.kind == "insert"]
    assert len(inserts) == 1
    assert inserts[0].after == "Second one here."


def test_a_deleted_paragraph_is_one_hunk():
    hunks = _hunks("First.\n\nGone.\n\nThird.", "First.\n\nThird.")

    deletes = [h for h in hunks if h.kind == "delete"]
    assert [h.before for h in deletes] == ["Gone."]


def test_two_unrelated_paragraphs_are_not_diffed_word_by_word():
    """Below the similarity threshold they are different paragraphs, not an edit."""
    hunks = _hunks("Alpha beta gamma delta.", "Completely unrelated sentence here.")

    assert len(hunks) == 1
    assert hunks[0].kind == "replace"


def test_recurring_corrections_share_a_signature_across_papers():
    """Digits are normalised so '5 kW'->'5~kW' and '7 kW'->'7~kW' count together."""
    one = _hunks("power of 5 kW", "power of 5~kW")[0]
    two = _hunks("power of 7 kW", "power of 7~kW")[0]

    assert one.signature == two.signature == "# kW -> #~kW"


# ---------------------------------------------------------------------------
# Word submissions
# ---------------------------------------------------------------------------

def _docx(path: Path, paragraphs: list[str], table: list[str] | None = None) -> Path:
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        grid = document.add_table(rows=1, cols=len(table))
        for cell, text in zip(grid.rows[0].cells, table):
            cell.text = text
    document.save(str(path))
    return path


def test_a_docx_is_read_as_paragraphs_including_its_tables(tmp_path: Path):
    """JACoW papers carry captions and references inside tables often enough
    that skipping them would lose real corrections."""
    path = _docx(tmp_path / "a.docx", ["First line.", "Second line."],
                 table=["In a cell."])

    assert word_paragraphs(path) == ["First line.", "Second line.", "In a cell."]


def test_word_pairs_diff_the_same_way_latex_ones_do(tmp_path: Path):
    before = _docx(tmp_path / "b.docx", ["The gun runs at 5 mhz."])
    after = _docx(tmp_path / "c.docx", ["The gun runs at 5 MHz."])

    hunks = diff_paragraphs(word_paragraphs(before), word_paragraphs(after),
                            paper="W", source="word")

    assert [(h.before, h.after) for h in hunks] == [("mhz.", "MHz.")]


# ---------------------------------------------------------------------------
# Whole papers, and failing safely
# ---------------------------------------------------------------------------

def _conference(tmp_path: Path) -> Path:
    def paper(code, orig, curr, tags=()):
        for role, text in (("original", orig), ("current", curr)):
            folder = tmp_path / code / role / "Source_Files"
            folder.mkdir(parents=True, exist_ok=True)
            (folder / f"{code}.tex").write_text(text, encoding="utf-8")
        return {"code": code, "friendly_id": 1, "contribution_id": 1, "title": code,
                "persons": [], "state": "accepted", "editor": "E",
                "tags": [{"code": t, "title": t} for t in tags],
                "revision_count": 2, "revisions": [], "note": ""}

    records = [
        paper("AAA01", "a beam of 10 MeV", "a beam of 10~MeV", tags=["TC12"]),
        paper("BBB02", "identical text", "identical text"),
    ]
    (tmp_path / "index.json").write_text(
        json.dumps({"event_id": 82, "papers": records}), encoding="utf-8")
    return tmp_path


def test_a_paper_carries_its_editor_tags_onto_its_diff(tmp_path: Path):
    """The tags are the labels; they have to travel with the hunks."""
    corpus = load(_conference(tmp_path))
    diff = diff_paper([p for p in corpus.papers if p.code == "AAA01"][0])

    assert diff.tags == ["TC12"]
    assert [(h.before, h.after) for h in diff.hunks] == [("10 MeV", "10~MeV")]


def test_an_unusable_paper_yields_its_reason_rather_than_an_empty_result(
        tmp_path: Path):
    corpus = load(_conference(tmp_path))
    diff = diff_paper([p for p in corpus.papers if p.code == "BBB02"][0])

    assert diff.hunks == []
    assert diff.note == "the paper file is unchanged between revisions"


def test_an_unreadable_file_is_noted_and_does_not_stop_the_run(tmp_path: Path):
    root = _conference(tmp_path)
    (root / "AAA01" / "current" / "Source_Files" / "AAA01.tex").write_bytes(
        b"\xff\xfe\x00binary")

    diffs = [diff_paper(p) for p in load(root).usable]

    assert len(diffs) == 1                # the run completed
    assert diffs[0].hunks or diffs[0].note


def test_the_diffs_are_written_one_file_per_paper_plus_a_combined_one(
        tmp_path: Path):
    corpus = load(_conference(tmp_path))
    diffs = [diff_paper(p) for p in corpus.usable]

    combined = write(diffs, tmp_path / "_diffs")

    assert (tmp_path / "_diffs" / "AAA01.json").exists()
    everything = json.loads(combined.read_text())
    assert everything[0]["hunks"][0]["after"] == "10~MeV"
    assert everything[0]["tags"] == ["TC12"]
