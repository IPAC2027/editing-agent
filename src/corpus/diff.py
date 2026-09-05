"""What the editors actually changed, one hunk at a time.

Step two of learning from a pulled conference. The pair on disk — the author's
revision and the editors' — is the ground truth; this module turns it into a
list of individual corrections that can be compared, one by one, against what
the agent would have proposed.

Three decisions shape it, and each exists to stop a particular kind of garbage.

**Paragraphs, not lines.** A LaTeX editor who rewraps a paragraph changes every
line in it without changing a word. Diffing lines would report that as dozens of
corrections and drown everything real. So both formats are normalised into
paragraphs first — blank-line separated for LaTeX, the document's own paragraphs
for Word — and only paragraphs that survive alignment are compared word by word.

**Word-level inside a paragraph.** ``10 MeV`` becoming ``10~MeV`` is one
correction of two tokens, not a rewritten paragraph. Reporting the whole
paragraph would make every hunk look large and make alignment in step three
hopeless.

**Whitespace-only differences are dropped, character changes are not.** The
distinction matters more here than it looks: a non-breaking space is a character
change and one of the most common real corrections in the corpus, while a line
rewrap or a trailing space is noise. The test suite pins both directions.
"""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from difflib import SequenceMatcher
from pathlib import Path

from src.corpus.index import PaperEntry

#: Below this ratio two paragraphs are different paragraphs, not an edit of one.
_PAIR_THRESHOLD = 0.55

#: Hunks longer than this are recorded but marked: they are rewrites, not
#: corrections, and treating them as single edits would distort every count.
_LARGE_HUNK_WORDS = 25


@dataclass
class Hunk:
    """One correction the editors made."""

    paper: str
    kind: str                 # replace | insert | delete
    before: str
    after: str
    unit: int                 # paragraph index in the original
    context: str = ""         # the paragraph it sits in, trimmed
    source: str = "latex"     # latex | word

    @property
    def words(self) -> int:
        return max(len(self.before.split()), len(self.after.split()))

    @property
    def large(self) -> bool:
        return self.words > _LARGE_HUNK_WORDS

    @property
    def signature(self) -> str:
        """A normalised form, for counting how often the same fix recurs."""
        return f"{_shape(self.before)} -> {_shape(self.after)}"


@dataclass
class PaperDiff:
    """Every correction on one paper, plus why there might be none."""

    paper: str
    source: str
    hunks: list[Hunk] = field(default_factory=list)
    note: str = ""
    tags: list[str] = field(default_factory=list)

    @property
    def small(self) -> list[Hunk]:
        return [h for h in self.hunks if not h.large]


# ---------------------------------------------------------------------------
# Reading the two formats into paragraphs
# ---------------------------------------------------------------------------

def latex_paragraphs(text: str) -> list[str]:
    """Blank-line separated blocks, with line wrapping normalised away.

    Wrapping is a presentation choice the author and the editor may each make
    differently; it is never the correction being studied.
    """
    blocks = re.split(r"\n[ \t]*\n+", text)
    return [" ".join(block.split()) for block in blocks if block.strip()]


def word_paragraphs(path: Path) -> list[str]:
    """A .docx as a list of paragraphs, tables included.

    Table cells carry references and captions in JACoW papers often enough that
    skipping them would lose real corrections.
    """
    from docx import Document

    document = Document(str(path))
    out = [" ".join(p.text.split()) for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(" ".join(p.text.split()) for p in cell.paragraphs)
    return [p for p in out if p]


# ---------------------------------------------------------------------------
# Diffing
# ---------------------------------------------------------------------------

#: Words, spaces and single symbols. Aligning on whole words alone loses the
#: corrections that *join* two words — ``250 pC`` becoming ``250~pC`` shifts
#: every later token and swallows whatever changed next to it. Splitting
#: punctuation out keeps the alignment local; the hunk is widened back to whole
#: words afterwards, so what gets reported is still readable.
_TOKEN_RE = re.compile(r"\w+|\s+|[^\w\s]")


def _tokens(paragraph: str) -> list[str]:
    return _TOKEN_RE.findall(paragraph)


def _no_space(tokens: list[str]) -> bool:
    return not any(t.isspace() for t in tokens)


def _group(opcodes, a: list[str], b: list[str]):
    """Merge changes that sit inside one word, then widen to word boundaries.

    Two passes. First, a change / short equal run / change sequence with no
    space in the middle is one correction, not three. Second, each change is
    extended outwards over the neighbouring non-space tokens so the hunk reads
    as ``250 pC -> 250~pC`` rather than `` -> ~``.
    """
    merged: list[list] = []
    for tag, i1, i2, j1, j2 in opcodes:
        if (tag != "equal" and merged and merged[-1][0] == "equal"
                and _no_space(a[merged[-1][1]:merged[-1][2]])
                and len(merged) >= 2 and merged[-2][0] != "equal"):
            gap = merged.pop()
            start = merged.pop()
            merged.append(["replace", start[1], i2, start[3], j2])
            continue
        merged.append([tag, i1, i2, j1, j2])

    out = []
    for index, (tag, i1, i2, j1, j2) in enumerate(merged):
        if tag == "equal":
            continue
        if index > 0 and merged[index - 1][0] == "equal":
            prev = merged[index - 1]
            while i1 > prev[1] and not a[i1 - 1].isspace():
                i1 -= 1
                j1 -= 1
        if index + 1 < len(merged) and merged[index + 1][0] == "equal":
            nxt = merged[index + 1]
            while i2 < nxt[2] and not a[i2].isspace():
                i2 += 1
                j2 += 1
        out.append((tag, max(i1, 0), i2, max(j1, 0), j2))
    return out


def _whitespace_only(before: str, after: str) -> bool:
    """True when nothing but the amount of ordinary space differs.

    ``10 MeV`` -> ``10~MeV`` is *not* whitespace-only: the tilde is a character,
    and that correction is the single most common one the agent makes.
    """
    return re.sub(r"[ \t\n]+", "", before) == re.sub(r"[ \t\n]+", "", after) and \
        before != after and not re.search(r"[~ ]", before + after)


def _shape(text: str) -> str:
    """A crude normalisation, so recurring corrections collapse together."""
    shaped = re.sub(r"\d+", "#", text.strip())
    shaped = re.sub(r"\s+", " ", shaped)
    return shaped[:60]


def diff_paragraph(before: str, after: str, *, paper: str, unit: int,
                   source: str) -> list[Hunk]:
    """Word-level corrections within one paragraph."""
    a, b = _tokens(before), _tokens(after)
    opcodes = SequenceMatcher(None, a, b, autojunk=False).get_opcodes()
    hunks: list[Hunk] = []
    for tag, i1, i2, j1, j2 in _group(opcodes, a, b):
        was, now = "".join(a[i1:i2]), "".join(b[j1:j2])
        if not was.strip() and not now.strip():
            continue
        if _whitespace_only(was, now):
            continue
        context = "".join(a[max(0, i1 - 20):i2 + 20]).strip()
        hunks.append(Hunk(paper=paper, kind="replace" if was and now else tag,
                          before=was.strip(), after=now.strip(),
                          unit=unit, context=context, source=source))
    return hunks


def diff_paragraphs(before: list[str], after: list[str], *, paper: str,
                    source: str) -> list[Hunk]:
    """Align paragraphs, then compare the ones that correspond.

    A paragraph that was inserted or deleted outright is recorded as one hunk
    rather than exploded into words: nothing useful is learned from diffing a
    new paragraph against nothing.
    """
    hunks: list[Hunk] = []
    matcher = SequenceMatcher(None, before, after, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag == "replace":
            olds, news = before[i1:i2], after[j1:j2]
            # Pair them up in order while they still resemble each other; the
            # remainder is a real insertion or deletion.
            for offset, (old, new) in enumerate(zip(olds, news)):
                ratio = SequenceMatcher(None, old, new).quick_ratio()
                if ratio >= _PAIR_THRESHOLD:
                    hunks.extend(diff_paragraph(old, new, paper=paper,
                                                unit=i1 + offset, source=source))
                else:
                    hunks.append(Hunk(paper=paper, kind="replace", before=old,
                                      after=new, unit=i1 + offset,
                                      context=old[:160], source=source))
            for offset, old in enumerate(olds[len(news):], start=len(news)):
                hunks.append(Hunk(paper=paper, kind="delete", before=old, after="",
                                  unit=i1 + offset, context=old[:160], source=source))
            for new in news[len(olds):]:
                hunks.append(Hunk(paper=paper, kind="insert", before="", after=new,
                                  unit=i2, context="", source=source))
        elif tag == "delete":
            for offset, old in enumerate(before[i1:i2]):
                hunks.append(Hunk(paper=paper, kind="delete", before=old, after="",
                                  unit=i1 + offset, context=old[:160], source=source))
        elif tag == "insert":
            for new in after[j1:j2]:
                hunks.append(Hunk(paper=paper, kind="insert", before="", after=new,
                                  unit=i1, context="", source=source))
    return hunks


def diff_paper(entry: PaperEntry) -> PaperDiff:
    """Every correction on one paper of a pulled conference."""
    usable, reason = entry.usability()
    result = PaperDiff(paper=entry.code, source=entry.kind, tags=list(entry.tags))
    if not usable:
        result.note = reason
        return result

    before_file = entry.original.main_file()
    after_file = entry.current.main_file()
    try:
        if entry.kind == "latex":
            before = latex_paragraphs(
                before_file.read_text(encoding="utf-8", errors="replace"))
            after = latex_paragraphs(
                after_file.read_text(encoding="utf-8", errors="replace"))
        elif entry.kind == "word":
            if before_file.suffix.lower() != ".docx":
                result.note = f"cannot read {before_file.suffix} — only .docx"
                return result
            before, after = word_paragraphs(before_file), word_paragraphs(after_file)
        else:
            result.note = f"no reader for {entry.kind}"
            return result
    except Exception as exc:  # noqa: BLE001 — one unreadable paper must not stop a run
        result.note = f"could not read: {type(exc).__name__}: {exc}"
        return result

    result.hunks = diff_paragraphs(before, after, paper=entry.code, source=entry.kind)
    if not result.hunks:
        result.note = "no textual difference (the change was elsewhere in the file)"
    return result


def write(diffs: list[PaperDiff], destination: Path) -> Path:
    """One json per paper, plus an all.json for whole-corpus analysis."""
    destination.mkdir(parents=True, exist_ok=True)
    for diff in diffs:
        (destination / f"{diff.paper}.json").write_text(
            json.dumps({"paper": diff.paper, "source": diff.source,
                        "tags": diff.tags, "note": diff.note,
                        "hunks": [asdict(h) for h in diff.hunks]},
                       indent=2, ensure_ascii=False),
            encoding="utf-8")
    combined = destination / "all.json"
    combined.write_text(json.dumps(
        [{"paper": d.paper, "source": d.source, "tags": d.tags, "note": d.note,
          "hunks": [asdict(h) for h in d.hunks]} for d in diffs],
        indent=2, ensure_ascii=False), encoding="utf-8")
    return combined
